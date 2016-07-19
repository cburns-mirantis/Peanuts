#!/usr/bin/env python3

# Mangement Interface
# Get DHCP Range
# Split document into different sections

# Post creation report or data output
# store report in a per run basis

import zipfile,time,argparse,requests,json,sys,os,paramiko
from requests.packages.urllib3.exceptions import InsecureRequestWarning
requests.packages.urllib3.disable_warnings(InsecureRequestWarning)
# Python's docx lib has an lxml dependency - do "pip3 install lxml" for Python3;
# if build fails, try doing "apt-get install libxml2-dev libxslt-dev" (possibly as root)
# and then re-run "pip3 install lxml"
from docx import Document
from docx.shared import Inches
from docx.enum.style import WD_STYLE
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_BREAK
# Replaces items in the docx
def docx_replace(old_file,new_file,rep):
    zin = zipfile.ZipFile (old_file, 'r')
    zout = zipfile.ZipFile (new_file, 'w')
    for item in zin.infolist():
        buffer = zin.read(item.filename)
        if (item.filename == 'word/document.xml'):
            res = buffer.decode("utf-8")
            for r in rep:
                res = res.replace(r,rep[r])
            buffer = res.encode("utf-8")
        zout.writestr(item, buffer)
    zout.close()
    zin.close()

# Get token from Keystone
def get_token():
    header = {"Content-Type": "application/json", 'accept': 'application/json'}
    creds = {"auth": {"tenantName": args.web_username,"passwordCredentials": {"username": args.web_username,"password": args.web_password}}}
    r = requests.post(url="https://" + args.host + ":8443/keystone/v2.0/tokens",headers=header,verify=False,json=creds)
    if r.status_code is not 200:
        sys.exit("Check Fuel username & password")
    return json.loads(r.text)['access']['token']['id']

def insert_table_at(table, paragraph):
    tbl = table._tbl
    p = paragraph._p
    p.addnext(tbl)
    paragraph.text = ""

def fuel_info():
    ssh = paramiko.SSHClient()
    fuel = {}
    ssh.set_missing_host_key_policy(paramiko.AutoAddPolicy())
    ssh.connect(args.host, username=args.ssh_username, password=args.ssh_password)
    ssh_stdin, ssh_stdout, ssh_stderr = ssh.exec_command("ip route ls | grep " + args.host + "| awk -e '{ print $3 }'")
    fuel['management_iface'] = ssh_stdout.readlines()[0].replace("\n","")
    ssh_stdin, ssh_stdout, ssh_stderr = ssh.exec_command("route -n | grep UG | awk -e '{ print $2 }'")
    fuel['gateway'] = ssh_stdout.readlines()[0].replace("\n","")
    # ssh_stdin, ssh_stdout, ssh_stderr = ssh.exec_command("fuel plugins list")
    # print(ssh_stdout.readlines())

    fuel['url'] = "https://" + args.host + ":8443"
    fuel['ssh'] = {"username":args.ssh_username,"password":args.ssh_password}
    fuel['web'] = {"username":args.web_username,"password":args.web_password}
    fuel['horizon'] = "172.16.0.2" # get horizon address
    return fuel

def gen_access_table(fuel, doc):
   row_count = 8
   col_count = 3
   #styles = doc.styles
   #style = styles.add_style('testStyle', WD_STYLE_TYPE.TABLE)

   table = doc.add_table(row_count, col_count)
   
   #table.style = 'LightGridAccent1'
   try:
      table.style = doc.styles['Grid Table 4 Accent 1']
   except:
      try:
         table.style = doc.styles['Table Grid']
      except:
         print("Could not apply style to tables")
   
   hdr_cells = table.rows[0].cells
   hdr_cells[0].text = 'Host'
   hdr_cells[1].text = 'IP address'
   hdr_cells[2].text = 'Comment(s)'

   line1 = table.rows[1].cells
   line1[0].text = 'Fuel UI Master Node URL'
   line1[1].text = fuel['url']

   line2 = table.rows[2].cells
   line2[0].text = 'Fuel UI Credentials'
   line2[1].text = fuel['web']['username'] + ' / ' + fuel['web']['password']

   line3 = table.rows[3].cells
   line3[0].text = 'Fuel Master Node IP'
   line3[1].text = args.host

   line4 = table.rows[4].cells
   line4[0].text = 'Fuel SSH Credentials (console)'
   line4[1].text = fuel['ssh']['username'] + ' / ' + fuel['ssh']['password']

   line5 = table.rows[5].cells
   line5[0].text = 'OpenStack Nodes SSH Credentials (console)'
   #line5[1].text = fuel['ssh']['username'] + ' / ' + fuel['ssh']['password']

   line6 = table.rows[6].cells
   line6[0].text = 'OpenStack Horizon URL'
   line6[1].text = fuel['horizon']

   line7 = table.rows[7].cells
   line7[0].text = 'OpenStack Credentials'
   line7[1].text = fuel['web']['username'] + ' / ' + fuel['web']['password']
   
   for p in doc.paragraphs:
      if "NODESHAREPLACE" in p.text:
         insert_table_at(table, p)
   return doc

# Get nodes from Fuel API
def get_nodes(token):
    header = {"X-Auth-Token": token,"Content-Type": "application/json"}
    return sorted(json.loads(requests.get(url="https://" + args.host + ":8443/api/nodes",headers=header, verify=False).text), key=lambda k: k['hostname'])

def gen_nodes_table(nodeData, doc):
   # Generate row_count by counting the number of nodes and adding 1 for the header row
   row_count = len(nodeData)+1
   col_count = 6

   # Create table for the data
   table = doc.add_table(row_count, col_count)
   try:
      table.style = doc.styles['Grid Table 4 Accent 1']
   except:
      try:
         table.style = doc.styles['Table Grid']
      except:
         print("Could not apply style to tables")

   # Generate the header row for the NodesVMS table
   hdr_cells = table.rows[0].cells
   hdr_cells[0].text = 'Hostname'
   hdr_cells[1].text = 'Role(s)'
   hdr_cells[2].text = 'Admin network IP address'
   hdr_cells[3].text = 'CPUxCores'
   hdr_cells[4].text = 'RAM'
   hdr_cells[5].text = 'HDD'

   # Modify each additional row using data for the node
   for nodeCounter, node in enumerate(nodeData):
      nodeRow = table.rows[nodeCounter+1].cells
      nodeRow[0].text = nodeData[nodeCounter]['hostname']
      nodeRow[1].text = {x+' ' for x in nodeData[nodeCounter]['roles']}
      nodeRow[2].text = nodeData[nodeCounter]['ip']
      nodeRow[3].text = str(nodeData[nodeCounter]['meta']['cpu']['total']) # Total or real?
      nodeRow[4].text = str(int(nodeData[nodeCounter]['meta']['memory']['total']/1048576)) + ' MB' # Total or max cap?
      nodeRow[5].text = [str(x['disk']) + ': ' + str(int(x['size']/1073741824)) + 'GB \n' for x in nodeData[nodeCounter]['meta']['disks']]

   for paragraph in doc.paragraphs:
      if "NODESVMSREPLACE" in paragraph.text:
         insert_table_at(table, paragraph)
   return doc

# Gathers network information on the cluster in question using the REST interface
def get_network_info(cluster_id):
    header = {"X-Auth-Token": token,"Content-Type": "application/json"}
    network_data = json.loads(requests.get(url="https://" + args.host + ":8443/api/clusters/" + str(cluster_id) + "/network_configuration/neutron/", headers=header, verify=False).text)
    networks = []
    for x in network_data['networks']:
        network = {}
        network['network_name'] = str(x['name']) if x['name'] is not None else 'no data'
        network['speed'] = 'no data'
        network['port_mode'] = 'no data' 
        network['ip_range'] = str(x['cidr']) if x['name'] is not None else 'no data'# Check this - might be ['networks'][x]['ip_ranges'] instead
        network['vlan'] = str(x['vlan_start']) if x['vlan_start'] is not None else 'no data'
        network['interface'] = 'no data'
        network['gateway'] = str(x['gateway']) if x['gateway'] is not None else 'no data'
        networks.append(network)
    return networks

# Creates the table for the Network Layout section  
def gen_network_layout_table(networkData, doc):
   row_count = len(networkData)+1
   col_count = 6

   table = doc.add_table(row_count, col_count)
   table.style = doc.styles['Grid Table 4 Accent 1']

   hdr_cells = table.rows[0].cells
   hdr_cells[0].text = 'Network name'
   hdr_cells[1].text = 'Speed'
   hdr_cells[2].text = 'Port mode'
   hdr_cells[3].text = 'IP Range'
   hdr_cells[4].text = 'VLAN'
   hdr_cells[5].text = 'Interface'

   for netCounter, network in enumerate(networkData):
      networkRow = table.rows[netCounter+1].cells
      networkRow[0].text = network['network_name']
      networkRow[1].text = network['speed']
      networkRow[2].text = network['port_mode']
      networkRow[3].text = network['ip_range']
      networkRow[4].text = network['vlan']
      networkRow[5].text = network['interface']

   for p in doc.paragraphs:
      if "NETWORKREPLACE" in p.text:
         insert_table_at(table, p)
   return doc
# Main 

# Handle Arguments
parser = argparse.ArgumentParser(description='Gather Fuel Screenshots')
parser.add_argument('-u', "--web-user", action="store", dest="web_username", type=str, help='Fuel Username',default="admin")
parser.add_argument('-p', "--web-pw", action="store", dest="web_password", type=str, help='Fuel Password',default="admin")
parser.add_argument('-su', "--ssh-user", action="store", dest="ssh_username", type=str, help='SSH Username',default="root")
parser.add_argument('-sp', "--ssh-pw", action="store", dest="ssh_password", type=str, help='SSH Password',default="r00tme")
parser.add_argument('-f', "--fuel", action="store", dest="host", type=str, help='Fuel FQDN or IP Ex. 10.20.0.2',required=True)
args = parser.parse_args()

if not os.path.exists("docs_full"):
   os.makedirs("docs_full")

# Generate the token for access:
try:
   token = get_token()
except:
   print("Could not generate access token - some doc builds may fail")

# Load the template document
doc = Document('templates_full/Start.docx')

for paragraph in doc.paragraphs:
   if "PAGEBREAK" in paragraph.text:
      paragraph.text = "\n"
      run = paragraph.add_run()
      run.add_break(WD_BREAK.PAGE)

doc = gen_access_table(fuel_info(), doc)
doc = gen_nodes_table(get_nodes(token), doc)
doc = gen_network_layout_table(get_network_info(1),doc)


doc.save("docs_full/testrun.docx")
'''
# 1.cover
try:
   entries = json.load(open("entries.json"))
   entries['DATE'] = time.strftime("%d %B, %Y")
   #docx_replace("templates_full/1.cover.docx","docs_full/1.cover.docx",entries)
   print("Built docs_full/1.cover.docx")
except:
   print("Failed to build docs_full/1.cover.docx")

# 2.intro

# print("Built docs_full/2.intro.docx")

# 3.architecture

# print("Built docs_full/3.architecture.docx")

# 4.network
# Some network info available on REST: /api/clusters/<cluster number>/network_configuration/
try:
   gen_network_layout_table(network_info(1))
   print("Built docs_full/4.network.docx")
except:
   print("Failed to build docs_full/4.network.docx")

# 5.nodes
gen_nodes_table(get_nodes(token))
print("Built docs_full/5.nodes.docx")

#   print("Failed to build docs_full/5.nodes.docx")

# 6.access
fuel = fuel_info()
gen_access_table(fuel)
access_replace = {
"FUELURL": fuel['url'],
"FUELIP": args.host,
"WEBUSER": fuel['web']['username'],
"WEBPW": fuel['web']['password'],
"HORURL": fuel['horizon'],
"SSHUSER": fuel['ssh']['username'],
"SSHUSER": fuel['ssh']['username'],
"SSHPW": fuel['ssh']['password']
}
docx_replace("templates_full/6.access.docx","docs_full/6.access.docx",access_replace)
print("Built docs_full/6.access.docx")
# 7.fuel
try:
   fuel_replace = {
   "TOTAL_NODES" : len(nodes) 
   }
   docx_replace("templates_full/7.fuel.docx","docs_full/7.fuel.docx",fuel_replace)
   print("Built docs_full/7.fuel.docx")
except:
   print("Failed to build docs_full/7.fuel.docx")
'''
