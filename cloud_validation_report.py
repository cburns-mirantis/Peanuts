#!/usr/bin/env python3
# Cloud Verification Report
# Version 0.01
# https://github.com/cburns-mirantis/Peanuts

import argparse,sys,requests,json,configparser,time,zipfile,paramiko,os,socket
from tabulate import tabulate
from docx import Document
from docx.shared import Inches
from collections import OrderedDict
from docx.enum.style import WD_STYLE
from requests.packages.urllib3.exceptions import InsecureRequestWarning
requests.packages.urllib3.disable_warnings(InsecureRequestWarning)
from novaclient.client import Client

parser = argparse.ArgumentParser(description='Gather Fuel Screenshots and generate a runbook for delivery to customer.')
parser.add_argument('-u', '--web-user', action='store', dest='web_username', type=str, help='Fuel Web Username',default='admin')
parser.add_argument('-p', '--web-pw', action='store', dest='web_password', type=str, help='Fuel Web Password',default='admin')
parser.add_argument('-hu', '--horizon-user', action='store', dest='horizon_username', type=str, help='Horizon Username',default='admin')
parser.add_argument('-hp', '--horizon-pw', action='store', dest='horizon_password', type=str, help='Horizon Password',default='admin')
parser.add_argument('-wp', '--web-port', action='store', dest='web_port', type=str, help='Fuel Web Port',default='8443')
parser.add_argument('-su', '--ssh-user', action='store', dest='ssh_username', type=str, help='Fuel SSH Username',default='root')
parser.add_argument('-sp', '--ssh-pw', action='store', dest='ssh_password', type=str, help='Fuel SSH Password',default='r00tme')
parser.add_argument('-f', '--fuel', action='store', dest='host', type=str, help='Fuel FQDN or IP Ex. 10.20.0.2',required=True)
args = parser.parse_args()

def get_token():
    header = {'Content-Type': 'application/json', 'accept': 'application/json'}
    creds = {'auth': {'tenantName': args.web_username,'passwordCredentials': {'username': args.web_username,'password': args.web_password}}}
    try:
        r = requests.post(url='https://' + args.host + ':' + args.web_port + '/keystone/v2.0/tokens',headers=header,verify=False,json=creds,timeout=15)
    except:
        sys.exit('Request timed out. Check Fuel web address and port.')
    if r.status_code is not 200:
        sys.exit('Check Fuel username & password')
    print('Successful auth.')
    return json.loads(r.text)['access']['token']['id']

def get_nodes(token):
    header = {'X-Auth-Token': token,'Content-Type': 'application/json'}
    return sorted(json.loads(requests.get(url='https://' + args.host + ':' + args.web_port + '/api/nodes',headers=header, verify=False).text), key=lambda k: k['hostname'])

def get_clusters(token):
    header = {'X-Auth-Token': token,'Content-Type': 'application/json'}
    return json.loads(requests.get(url='https://' + args.host + ':' + args.web_port + '/api/clusters',headers=header, verify=False).text)

def get_network(token,cluster_id):
    header = {'X-Auth-Token': token,'Content-Type': 'application/json'}
    return json.loads(requests.get(url='https://' + args.host + ':' + args.web_port + '/api/clusters/' + str(cluster_id) + '/network_configuration/neutron/', headers=header, verify=False).text)

def count_cluster_nodes(id):
    count = 0
    for n in nodedata:
        if str(id) in str(n['cluster']):
            count += 1
    return str(count)

def gen_deployment_overview(cluster):
    row_count = 6
    col_count = 2

    heading = runbook.add_heading('Deployment Overview',level=1)
    heading.alignment = 1
    table = runbook.add_table(row_count, col_count)
    table.style = runbook.styles['Light Grid Accent 1']

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = ''
    hdr_cells[1].text = ''

    line1 = table.rows[1].cells
    line1[0].text = 'Cloud Purpose'
    line1[1].text = entries['DEPLOYMENT']['CLOUD_PURPOSE']
    line1 = table.rows[2].cells
    line1[0].text = 'Mirantis Openstack Version'
    line1[1].text = cluster['fuel_version']
    line1 = table.rows[3].cells
    line1[0].text = 'Host OS'
    line1[1].text = nodedata[0]['os_platform']
    line1 = table.rows[4].cells
    line1[0].text = 'Number of Nodes'
    line1[1].text = count_cluster_nodes(cluster['id'])
    line1 = table.rows[5].cells
    line1[0].text = 'Required Openstack configuration in high level'
    line1[1].text = "TBD"

def docx_replace(old_file,new_file,rep):
    zin = zipfile.ZipFile (old_file, 'r')
    zout = zipfile.ZipFile (new_file, 'w')
    for item in zin.infolist():
        buffer = zin.read(item.filename)
        if (item.filename == 'word/document.xml'):
            res = buffer.decode('utf-8')
            for r in rep:
                res = res.replace(r,rep[r])
            buffer = res.encode('utf-8')
        zout.writestr(item, buffer)
    zout.close()
    zin.close()

def launch_5_vms_and_verify():
    try:
        keypair = nova_client.keypairs.create(name="HA Testing")
        secgroup = nova_client.security_groups.create(name="HA Testing",description="HA Testing")
        nova_client.security_group_rules.create(secgroup.id, ip_protocol="tcp",from_port=22, to_port=22)

        image = nova_client.images.find(name="TestVM")
        flavor = nova_client.flavors.find(name="m1.micro")
        net = nova_client.networks.find(label="admin_internal_net")
    # except nova_client.exceptions.Unauthorized:
        # sys.exit("Horizon authentication failed. Check username & password.")
    except (requests.exceptions.ConnectionError, requests.exceptions.ChunkedEncodingError, BadRequest,ClientException) as e:
        print("test failed")
        return
    floating_ips = []
    servers = []
    for i,s in enumerate(range(2)):
        try:
            floating_ips.append( nova_client.floating_ips.create(nova_client.floating_ip_pools.list()[0].name))
            servers.append( nova_client.servers.create(name="HA Test-" + str(i), image=image,flavor=flavor, key_name="HA Testing", nics=[{"net-id": net.id}]))
        except (requests.exceptions.ConnectionError, requests.exceptions.ChunkedEncodingError, BadRequest,ClientException) as e:
            print("test failed")
            return
    while(True):
        for s in nova_client.servers.list():
            if "HA Test-" in s.name and "ACTIVE" not in s.status:
                print("Hosts still booting, waiting 5 seconds...")
                time.sleep(5)
                continue
            break
        break

    print("Hosts booted. Attaching floating ips.")
    for i,s in enumerate(servers):
        if "HA Test-" in s.name:
            try:
                s.add_floating_ip(floating_ips[i])
            except (requests.exceptions.ConnectionError, requests.exceptions.ChunkedEncodingError, BadRequest,ClientException) as e:
                print("test failed")
                return

    for f in floating_ips:
        while(socket.socket(socket.AF_INET, socket.SOCK_STREAM).connect_ex((f.ip, 22)) is not 0):
            print("waiting")
            time.sleep(5)
        ssh = paramiko.SSHClient()
        ssh.set_missing_host_key_policy(paramiko.AutoAddPolicy())
        ssh.connect(f.ip, username="cirros", password="cubswin:)")
        if ssh.get_transport().is_active():
            print(f.ip,"has ssh")

    nova_client.keypairs.delete(keypair)
    nova_client.security_groups.delete(secgroup)
    print("destroying servers")
    for s in servers:
        s.delete()

    while(len(nova_client.servers.list()) is not 0):
        print("waiting for servers to finish destroying...")
        time.sleep(5)

    print("releasing floating ips...")
    for f in floating_ips:
        f.delete()

    # while(len(nova_client.floating_ips.list())):
    #     time.sleep(5)

def perform_ha_tests(id):
    hosts = []
    # print("Starting HA Testing...")
    for n in nodedata:
        if str(id) in str(n['cluster']) and 'controller' in n['roles']:
            hosts.append([n['id'],n['name'],n['ip']])
    print(tabulate(hosts,headers=["ID","Name","IP"]))

    node_id = input("Enter the ID of the controller to shutdown:\n")

    for n in nodedata:
        if str(node_id) in str(n['id']):
            node_ip = n['ip']

    ssh = paramiko.SSHClient()
    ssh.set_missing_host_key_policy(paramiko.AutoAddPolicy())
    try:
        ssh.connect(args.host, username=args.ssh_username, password=args.ssh_password)
    except paramiko.ssh_exception.AuthenticationException:
        sys.exit("SSH authentication failed. Check username & password.")
    print('Sending shutdown command...')
    ssh_stdin, ssh_stdout, ssh_stderr = ssh.exec_command('ssh root@' + args.host + ' -t \'ssh root@' + node_ip + ' shutdown -h 0\'')

    while(True):
        if os.system('ping -c 1 -t 5 ' + node_ip) == 0:
            print("Waiting 15 seconds for host to become unavailable...")
            time.sleep(15)
            continue
        else:
            print("Host is down.")
            break
    launch_5_vms_and_verify()
    input("press enter once host has booted.")
    ssh_stdin, ssh_stdout, ssh_stderr = ssh.exec_command('ssh root@' + args.host + ' -t \'ssh root@' + node_ip + ' echo c > /proc/sysrq-trigger\'')
    launch_5_vms_and_verify()
    input("press enter once host has booted.")
    # ssh_stdin, ssh_stdout, ssh_stderr = ssh.exec_command('ssh root@' + args.host + ' -t \'ssh root@' + node_ip + ' ip link set dev eth0 down\'')
    # launch_5_vms_and_verify()
    # input("press enter once host has booted.")



def gen_ha(results):
    row_count = 6
    col_count = 2


    heading = runbook.add_heading('HA Gracefult Shutdown Test Results',level=1)
    heading.alignment = 1
    table = runbook.add_table(row_count, col_count)
    table.style = runbook.styles['Light Grid Accent 1']
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Test'
    hdr_cells[1].text = 'Result'
    line1 = table.rows[1].cells
    line1[0].text = 'Gracefully shutdown one of the controller nodes.'
    line1[1].text = "TBD"
    line1 = table.rows[2].cells
    line1[0].text = 'Boot 5 VMs and verify that they all are up and running.'
    line1[1].text = "TBD"
    line1 = table.rows[3].cells
    line1[0].text = 'Turn on controller and wait until all services are up (Time limit 30 min)'
    line1[1].text = "TBD"
    line1 = table.rows[4].cells
    line1[0].text = 'Boot 5 VMs and verify that they all are up and running.'
    line1[1].text = "TBD"
    line1 = table.rows[5].cells
    line1[0].text = 'Return the cluster to initial state.'
    line1[1].text = "TBD"

    runbook.add_page_break()
    heading = runbook.add_heading('HA Hard Power Off Test Results',level=1)
    heading.alignment = 1
    table = runbook.add_table(row_count, col_count)
    table.style = runbook.styles['Light Grid Accent 1']
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Test'
    hdr_cells[1].text = 'Result'
    line1 = table.rows[1].cells
    line1[0].text = 'Hard power off one of the controller nodes.'
    line1[1].text = "TBD"
    line1 = table.rows[2].cells
    line1[0].text = 'Boot 5 VMs and verify that they all are up and running.'
    line1[1].text = "TBD"
    line1 = table.rows[3].cells
    line1[0].text = 'Turn on controller and wait until all services are up, but not longer than 30 min.'
    line1[1].text = "TBD"
    line1 = table.rows[4].cells
    line1[0].text = 'Boot 5 VMs and verify that they all are up and running.'
    line1[1].text = "TBD"
    line1 = table.rows[5].cells
    line1[0].text = 'Return the cluster to initial state.'
    line1[1].text = "TBD"

    runbook.add_page_break()
    heading = runbook.add_heading('HA Network Cut Test Results',level=1)
    heading.alignment = 1
    table = runbook.add_table(row_count, col_count)
    table.style = runbook.styles['Light Grid Accent 1']
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Test'
    hdr_cells[1].text = 'Result'
    line1 = table.rows[1].cells
    line1[0].text = 'Cut network communication on one of the controller nodes.'
    line1[1].text = "TBD"
    line1 = table.rows[2].cells
    line1[0].text = 'Boot 5 VMs and verify that they all are up and running.'
    line1[1].text = "TBD"
    line1 = table.rows[3].cells
    line1[0].text = 'Turn on controller and wait until all services are up (Time limit 30 min)'
    line1[1].text = "TBD"
    line1 = table.rows[4].cells
    line1[0].text = 'Boot 5 VMs and verify that they all are up and running.'
    line1[1].text = "TBD"
    line1 = table.rows[5].cells
    line1[0].text = 'Return the cluster to initial state.'
    line1[1].text = "TBD"


token = get_token()
nodedata = get_nodes(token)
clusters = get_clusters(token)
entries = configparser.ConfigParser()
entries.read('entries.cfg')

table = []
# choose environment
for c in clusters:
    row = [c['id'],c['name'],c['fuel_version'],c['status'],count_cluster_nodes(c['id'])]
    table.append(row)

print (tabulate(table,headers=["ID","Name","Version","Status","Nodes"]))

env_id = input("\nEnter the ID of the environment that you would like to run the verifcation report:\n")

networks = get_network(token,env_id)

credentials = {
    "version": "2",
    "username": args.horizon_username,
    "api_key": args.horizon_password,
    "auth_url": "http://" + networks['public_vip'] + ":5000/v2.0",
    "project_id": args.horizon_username
}

nova_client = Client(**credentials)




# print("Generating cover page...")
replaces = {
"CUSTOMER": entries['COVER']['CUSTOMER'],
"ENV": entries['COVER']['ENV'],
"RELEASE": entries['COVER']['RELEASE'],
"DATE": time.strftime('%d %B, %Y'),
"AUTHORS": entries['COVER']['AUTHORS']
}
docx_replace('template.docx','cover.docx',replaces)

runbook = Document('cover.docx')

for c in clusters:
    if env_id in str(c['id']):
        # print("generating report for",c['name'])
        gen_deployment_overview(c)
        runbook.add_page_break()
        gen_ha(perform_ha_tests(c['id']))
        break
    else:
        continue
    sys.exit('Environment not found.')

runbook.save('Cloud Verification.docx')
os.remove('cover.docx')
