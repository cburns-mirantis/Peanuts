#!/usr/bin/env python3
# Version 0.03
# https://github.com/cburns-mirantis/Peanuts

import zipfile,time,argparse,requests,json,sys,os,paramiko,shutil,math,re,configparser,select,threading
from tabulate import tabulate
from requests.packages.urllib3.exceptions import InsecureRequestWarning
requests.packages.urllib3.disable_warnings(InsecureRequestWarning)

from PIL import Image
from docx import Document
from docx.shared import Inches
from collections import OrderedDict
from docx.enum.style import WD_STYLE
import nwdiag.parser
from nwdiag.drawer import DiagramDraw
from nwdiag.builder import ScreenNodeBuilder
from selenium import webdriver
from selenium.webdriver.common.keys import Keys
from selenium.common.exceptions import NoSuchElementException
import socketserver as SocketServer

# Handle Arguments
parser = argparse.ArgumentParser()
parser.add_argument('--web-user', action='store', dest='web_username', type=str)
parser.add_argument('--web-pw', action='store', dest='web_password', type=str)
parser.add_argument('--horizon-user', action='store', dest='horizon_username', type=str)
parser.add_argument('--horizon-pw', action='store', dest='horizon_password', type=str)
parser.add_argument('--fuel-port', action='store', dest='ssh_port', type=str)
parser.add_argument('--ssh-user', action='store', dest='ssh_username', type=str)
parser.add_argument('--ssh-pw', action='store', dest='ssh_password', type=str)
parser.add_argument('--environment', action='store', dest='environment', type=str)
parser.add_argument('--customer-name', action='store', dest='customer_name', type=str)
parser.add_argument('--time-zone', action='store', dest='timezone', type=str)
parser.add_argument('--entitlement', action='store', dest='entitlement', type=str)
parser.add_argument('--environment-type', action='store', dest='environment_type', type=str)
parser.add_argument('--customer-manager', action='store', dest='customer_manager', type=str)
parser.add_argument('--deployment-engineer', action='store', dest='deployment_engineer', type=str)
parser.add_argument('--filename', action='store', dest='filename', type=str)
parser.add_argument('--fuel-address', action='store', dest='host', type=str)
parser.add_argument('--test',action="append",dest="tests", type=str)
args = parser.parse_args()

# Get token from Keystone
def get_token():
    header = {'Content-Type': 'application/json', 'accept': 'application/json'}
    creds = {'auth': {'tenantName': args.web_username,'passwordCredentials': {'username': args.web_username,'password': args.web_password}}}
    try:
        r = requests.post(url='https://localhost:8443/keystone/v2.0/tokens',headers=header,verify=False,json=creds,timeout=15)
    except:
        sys.exit(12)
    if r.status_code is not 200:
        sys.exit(13)
    return json.loads(r.text)['access']['token']['id']

def get_nodes(token):
    header = {'X-Auth-Token': token,'Content-Type': 'application/json'}
    return sorted(json.loads(requests.get(url='https://localhost:8443/api/nodes',headers=header, verify=False).text), key=lambda k: k['hostname'])

def get_cluster(token,cluster_id):
    header = {'X-Auth-Token': token,'Content-Type': 'application/json'}
    cluster = json.loads(requests.get(url='https://localhost:8443/api/clusters/' + cluster_id,headers=header, verify=False).text)
    try:
        if 'Cluster not found' in cluster['message']:
            sys.exit(11)
    except KeyError:
        pass
    if 'operational' not in cluster['status']:
        sys.exit(10)
    return cluster

def get_version(token):
    header = {'X-Auth-Token': token,'Content-Type': 'application/json'}
    return json.loads(requests.get(url='https://localhost:8443/api/version',headers=header, verify=False).text)['release']

def get_network(token,cluster_id):
    header = {'X-Auth-Token': token,'Content-Type': 'application/json'}
    return json.loads(requests.get(url='https://localhost:8443/api/clusters/' + str(cluster_id) + '/network_configuration/neutron/', headers=header, verify=False).text)

def get_test_result(token,cluster_id):
    header = {'X-Auth-Token': token,'Content-Type': 'application/json'}
    return json.loads(requests.get(url='http://localhost:8777/v1/testruns/last/' + str(cluster_id), headers=header, verify=False).text)

def start_ostf(token,cluster_id,username,password):
    header = {'X-Auth-Token': token,'Content-Type': 'application/json'}
    tests = []
    for t in args.tests:
        tests.append({
            "testset": t,
            "tests": [],
            "metadata": {
                "cluster_id": cluster_id,
                "ostf_os_access_creds": {
                        "ostf_os_username": username,
                        "ostf_os_password": password,
                        "ostf_os_tenant_name": username
                }
            }
        })

    r = requests.post(url='http://localhost:8777/v1/testruns/',headers=header,verify=False,json=tests,timeout=15)
    if r.status_code is not 200:
        sys.exit(14)

def wait_and_collect_ostf(token,cluster_id):
    while True:
        tests_completed = 0
        ostf_results = get_test_result(token,cluster_id)
        for r in ostf_results:
            if 'finished' in r['status']:
                tests_completed += 1
        if tests_completed == 6:
            return ostf_results
            break
        tests_completed = 0
        time.sleep(15)

def gen_ostf_table(tests,environment_name):
    for test_count,t in enumerate(tests):
        row_count = int(len(t['tests']))+1
        heading = runbook.add_heading('Environment: ' + environment_name + ' ' + t['testset'].title() + ' Health Check',level=1)
        heading.alignment = 1
        table = runbook.add_table(row_count, 4)
        table.style = runbook.styles['Light Grid Accent 1']
        table.autofit = False

        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Test'
        hdr_cells[1].text = 'Result'
        hdr_cells[2].text = 'Time Taken'
        hdr_cells[3].text = 'Message'

        for row_count,r in enumerate(t['tests']):
            line = table.rows[row_count+1].cells
            line[0].text = r['name']
            line[1].text = r['status']
            line[2].text = "%.2f" % r['taken']
            line[3].text = r['message']

        runbook.add_page_break()
        if test_count+1 is len(tests):
            runbook.add_page_break()

# Get information about the Fuel Instance through SSH
def fuel_info():
    ssh = paramiko.SSHClient()
    fuel = {}
    ssh.set_missing_host_key_policy(paramiko.AutoAddPolicy())
    ssh.connect(args.host, username=args.ssh_username, password=args.ssh_password,port=int(args.ssh_port))
    ssh_stdin, ssh_stdout, ssh_stderr = ssh.exec_command('ip route ls | grep ' + args.host + '| awk \'{ print $3 }\'')

    fuel['management_iface'] = ssh_stdout.readlines()[0].replace('\n','')
    ssh_stdin, ssh_stdout, ssh_stderr = ssh.exec_command('route -n | grep UG | awk \'{ print $2 }\'')
    # if ssh_stdout.readlines()[0] is not None:
    fuel['gateway'] = ssh_stdout.readlines()[0].replace('\n','')
    # else:
        # fuel['gateway'] = "No Gateway"
    ssh_stdin, ssh_stdout, ssh_stderr = ssh.exec_command('cat /etc/fuel/fuel-uuid')
    fuel['UUID'] = ssh_stdout.readlines()[0].replace('\n','')

    # fuel['url'] = 'https://' + args.host + ':8443'
    # fuel['ssh'] = {'username':args.ssh_username,'password':args.ssh_password}
    # fuel['web'] = {'username':args.web_username,'password':args.web_password}
    return fuel

# Generate 'Access Information' table
def gen_access_table():
   row_count = 8
   col_count = 2

   heading = runbook.add_heading('Access Information',level=1)
   heading.alignment = 1
   table = runbook.add_table(row_count, col_count)
   table.style = runbook.styles['Light Grid Accent 1']

   hdr_cells = table.rows[0].cells
   hdr_cells[0].text = ''
   hdr_cells[1].text = ''
   # hdr_cells[2].text = 'Comment(s)'

   line1 = table.rows[1].cells
   line1[0].text = 'Fuel UI Master Node URL'
   line1[1].text = 'https://' + args.host + ':8443'

   line2 = table.rows[2].cells
   line2[0].text = 'Fuel UI Credentials'
   line2[1].text = args.web_username + ' / ' + args.web_password

   line3 = table.rows[3].cells
   line3[0].text = 'Fuel Master Node IP'
   line3[1].text = args.host

   line4 = table.rows[4].cells
   line4[0].text = 'Fuel SSH Credentials'
   line4[1].text = args.ssh_username + ' / ' + args.ssh_password

   line5 = table.rows[5].cells
   line5[0].text = 'OpenStack Nodes SSH Credentials'
   line5[1].text = args.ssh_username + ' / ' + args.ssh_password

   line6 = table.rows[6].cells
   line6[0].text = 'OpenStack Horizon URL'
   # line6[1].text = networkdata['vip']

   line7 = table.rows[7].cells
   line7[0].text = 'OpenStack Horizon Credentials'
   line7[1].text = args.horizon_username + ' / ' + args.horizon_password

# Handle entitlement table based on Support entitlement level
def entitlement_handler(entitlement):
    support_services = {}
    if int(entitlement) is 1:
        support_services =  {
            'ENTITLEMENT':'8 x 5',
            'E_DAYS':'Monday - Friday',
            'E_HOURS':'9am - 5pm',
            'E_SEV1':'4 Business Hours',
            'E_SEV2':'8 Business Hours',
            'E_SEV3':'24 Business Hours',
            'E_SEV4':'48 Business Hours',
        }
    if int(entitlement) is 2:
        support_services =  {
            'ENTITLEMENT':'24 x 7',
            'E_DAYS':'24 x 7',
            'E_HOURS':'24 x 7',
            'E_SEV1':'1 Business Hour',
            'E_SEV2':'2 Business Hours',
            'E_SEV3':'4 Business Hours',
            'E_SEV4':'8 Business Hours',
        }
    if int(entitlement) is 3:
        support_services =  {
            'ENTITLEMENT':'24 x 7',
            'E_DAYS':'24 x 7',
            'E_HOURS':'24 x 7',
            'E_SEV1':'15 Minutes',
            'E_SEV2':'1 Business Hour',
            'E_SEV3':'4 Business Hours',
            'E_SEV4':'8 Business Hours',
        }
    return support_services

# Generate 'Support Information' table
def gen_support_table():
    row_count = 14
    col_count = 2
    if args.entitlement is 3:
        row_count += 1
    entitlements = entitlement_handler(args.entitlement)

    heading = runbook.add_heading('Support Information',level=1)
    heading.alignment = 1
    table = runbook.add_table(row_count, col_count)
    table.style = runbook.styles['Light Grid Accent 1']

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = ''
    hdr_cells[1].text = ''

    line1 = table.rows[1].cells
    line1[0].text = 'Customer Name'
    line1[1].text = args.customer_name

    line2 = table.rows[2].cells
    line2[0].text = 'Environment ID'
    line2[1].text = fuel['UUID']

    line3 = table.rows[3].cells
    line3[0].text = 'Support Entitlement Package'
    line3[1].text = entitlements['ENTITLEMENT']

    line4 = table.rows[4].cells
    line4[0].text = 'Days of Direct Support'
    line4[1].text = entitlements['E_DAYS']

    line5 = table.rows[5].cells
    line5[0].text = 'Hours of Direct Support'
    line5[1].text = entitlements['E_HOURS']

    line6 = table.rows[6].cells
    line6[0].text = 'Support Timezone'
    line6[1].text = args.timezone

    line7 = table.rows[7].cells
    line7[0].text = 'Support Phone Numbers'
    line7[1].text = '+1 (650) 963-9828, +1 (925) 808-FUEL'

    line8 = table.rows[8].cells
    line8[0].text = 'Support Email'
    line8[1].text = 'support@mirantis.com'

    line9 = table.rows[9].cells
    line9[0].text = 'Support Website'
    line9[1].text = 'https://support.mirantis.com/'

    line10 = table.rows[10].cells
    line10[0].text = 'Severity 1 SLA'
    line10[1].text = entitlements['E_SEV1']

    line11 = table.rows[11].cells
    line11[0].text = 'Severity 2 SLA'
    line11[1].text = entitlements['E_SEV2']

    line12 = table.rows[12].cells
    line12[0].text = 'Severity 3 SLA'
    line12[1].text = entitlements['E_SEV3']

    line13 = table.rows[13].cells
    line13[0].text = 'Severity 4 SLA'
    line13[1].text = entitlements['E_SEV4']

    if args.entitlement is 3:
       line14 = table.rows[14].cells
       line14[0].text = 'Customer Success Manager'
       line14[1].text = args.customer_manager

# Replaces items in the docx
def docx_replace(old_file,new_file,rep):
    zin = zipfile.ZipFile (old_file, 'r')
    zout = zipfile.ZipFile (new_file, 'w')
    for item in zin.infolist():
        buffer = zin.read(item.filename)
        if (item.filename == 'word/document.xml'):
            res = buffer.decode('utf-8')
            for r in rep:
                res = res.replace(r,rep[r])
            buffer = res.encode('utf-8')
        zout.writestr(item, buffer)
    zout.close()
    zin.close()

# Generate 'Nodes' table
def gen_nodes_table(cluster_id):

    nodes = []
    col_count = 6
    for node in nodedata:
        if node['cluster'] is int(cluster_id):
            nodes.append(node)

    heading = runbook.add_heading('Nodes',level=1)
    heading.alignment = 1

    table = runbook.add_table(len(nodes)+1, col_count)
    table.style = runbook.styles['Light Grid Accent 1']

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Hostname'
    hdr_cells[1].text = 'Role(s)'
    hdr_cells[2].text = 'Admin network IP address'
    hdr_cells[3].text = 'Cores x CPU'
    hdr_cells[4].text = 'RAM'
    hdr_cells[5].text = 'HDD'

    for nodeCounter, node in enumerate(nodes,1):
        nodeRow = table.rows[nodeCounter].cells
        nodeRow[0].text = node['hostname']
        nodeRow[1].text = [(x+', ' if x in node['roles'][:-1] else x) for x in node['roles']]
        nodeRow[2].text = node['ip']
        nodeRow[3].text = str(node['meta']['cpu']['total']) + ' x ' + str(node['meta']['cpu']['spec'][0]['model']) # Total or real?
        nodeRow[4].text = str(int(node['meta']['memory']['total']/1048576)) + ' MB'
        nodeRow[5].text = [str(x['name']) + ': ' + str(int(x['size']/1073741824)) + 'GB     ' for x in node['meta']['disks']]

def get_network_layout():
    networks = []
    for x in networkdata['networks']:
        network = {}
        network['network_name'] = str(x['name']) if x['name'] is not None else '(no data)'
        for i in nodedata:
            if len(i['network_data']):
                for z in i['network_data']:
                        for y in i['meta']['interfaces']:
                            if y['name'] == z['dev']:
                                network['speed'] = str(y['max_speed']) if y['max_speed'] is not None else '(no data)'
                                break

        network['ip_range'] = str(x['cidr']) if x['name'] is not None else '(no data)'# Check this - might be ['networks'][x]['ip_ranges'] instead
        network['vlan'] = str(x['vlan_start']) if x['vlan_start'] is not None else 'No VLAN'
        # network['gateway'] = str(x['gateway']) if x['gateway'] is not None else '(no data)'
        networks.append(network)
    return networks

# Generate 'Network Layout' table
def gen_network_layout_table(networks,env):
    row_count = len(networks)+1
    col_count = 4

    heading = runbook.add_heading('Network Layout - Environment ' + env,level=1)
    heading.alignment = 1

    table = runbook.add_table(row_count, col_count)
    table.style = runbook.styles['Light Grid Accent 1']

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Network name'
    hdr_cells[1].text = 'Speed'
    hdr_cells[2].text = 'IP Range'
    hdr_cells[3].text = 'VLAN'

    for netCounter, network in enumerate(networks):
        networkRow = table.rows[netCounter+1].cells
        networkRow[0].text = network['network_name']
        # networkRow[1].text = network['speed']
        networkRow[2].text = network['ip_range']
        networkRow[3].text = network['vlan']

def wait_for_page_tag_name(name):
    while(True):
        try:
            test = driver.find_element_by_name(name)
        except NoSuchElementException:
            continue
        break

def wait_for_page_tag_class(class_name):
    while(True):
        try:
            test = driver.find_element_by_class_name(class_name)
        except NoSuchElementException:
            continue
        break

# Handle webpage screenshots
def screenshot(page,name,tag,fix=False):
    if fix:
        driver.get('https://localhost:8443')
        time.sleep(1)
    if page is not None:
        driver.get(page)
    if tag is not None:
        wait_for_page_tag_class(tag)
    total_height = driver.execute_script('return document.body.parentNode.scrollHeight')

    viewport_height = driver.execute_script('return window.innerHeight')
    if total_height / viewport_height > 1 and total_height - viewport_height > 100 :
        passes = math.ceil(total_height / viewport_height)
        for i in range(passes):
            if i == 0:
                driver.save_screenshot('screens/' + name + '_0.png')
                add_picture_page('screens/' + name + '_0.png')
                continue
            driver.execute_script('window.scrollTo(0, '+ str(i*viewport_height) + ');')
            time.sleep(.2)
            driver.save_screenshot('screens/' + name + '_' + str(i) + '.png')
            add_picture_page('screens/' + name + '_' + str(i) + '.png')
    elif total_height - viewport_height < 100:
        driver.execute_script('window.scrollTo(0, '+ str(viewport_height) + ');')
        time.sleep(.2)
        driver.save_screenshot('screens/' + name + '.png')
        add_picture_page('screens/' + name + '.png')
    else:
        driver.save_screenshot('screens/' + name + '.png')
        add_picture_page('screens/' + name + '.png')

# Checks PATH for chromedriver dependency
def cmd_exists(cmd):
    return any(
        os.access(os.path.join(path, cmd), os.X_OK)
        for path in os.environ['PATH'].split(os.pathsep)
    )

def generate_server_string(token, cluster, host, physical_NICs):
    # Create server string
    server_string = str(host['roles']).replace("[","").replace("]","").replace("'","") + ","

    for NIC in physical_NICs:
        server_string += NIC  + ","
    return server_string

def get_node_NIC_hardware(token, hosts, cluster, net_name):
    nic = ""
    new_nodes = []
    known_roles = []

    a = []

    # for host in cluster:
    # Pre-count the total number of nodes with the same type of hardware
    physical_NICs = []
    for host in hosts:

        if host not in known_roles:
            net_count = 0
            physical_NICs = []
            tmp_NIC = ""
            for networks in host["network_data"]:
                if host["network_data"][net_count]['dev'] not in physical_NICs:
                    physical_NICs.append(host["network_data"][net_count]['dev'])
        new_node = generate_server_string(token, cluster, host, physical_NICs).replace(",","")

        node_found = False
        i = 0
        while i < len(a):
            if a[i][0] == new_node:
                node_found = True
            i += 1

        if node_found:
            i = 0
            while i < len(a):
                if a[i][0] == new_node:
                     a[i][1] += 1
                i += 1

        else:
            node = []
            node.append(new_node)
            node.append(1)
            a.append(node)

    for host in hosts:
        if host not in known_roles:
            net_count = 0
            physical_NICs = []
            tmp_NIC = ""

            title = ""
            i = 0
            while i < len(a):
                node_name = generate_server_string(token, cluster, host, physical_NICs).replace(",","")
                if a[i][0] == node_name:
                    title = str(a[i][0]) + " " + str(a[i][1])
                i += 1

            for networks in host["network_data"]:
                if host["network_data"][net_count]['dev'] not in physical_NICs:
                    physical_NICs.append(host["network_data"][net_count]['dev'])

                i = 0
                while i < len(a):
                    node_name = generate_server_string(token, cluster, host, physical_NICs).replace(",","")
                    if a[i][0] == node_name:
                        title = str(host['roles']) + "x" + str(a[i][1])
                    i += 1

                try:
                    if host["network_data"][net_count]['name'] == net_name:

                        if host["network_data"][net_count]['vlan'] is not None:
                            tmp_NIC += '"' + title + '" [address = "'+ host["network_data"][net_count]['dev'] + ', ' + host["network_data"][net_count]['ip'] + ', VLAN ' + str(host["network_data"][net_count]['vlan']) +'"];'
                        else:
                            tmp_NIC += '"' + title + '" [address = "'+ host["network_data"][net_count]['dev'] + ', ' + host["network_data"][net_count]['ip'] + '"];'
                except:
                    tmp_NIC += '"' + title + '" [address = "'+ host["network_data"][net_count]['dev'] + '"];'
                    pass
                net_count += 1

            new_node = generate_server_string(token, cluster, host, physical_NICs).replace(",","")

            if new_node not in new_nodes and new_node != ",":
                #if new_node == target_node:
                new_nodes.append(generate_server_string(token, cluster, host, physical_NICs).replace(",",""))
                nic += tmp_NIC
    return nic

def create_network_diagram(token, networks, cluster):
    hosts = get_nodes(token)
    diagram_input = "nwdiag {"
    for network in networks:
        try:
            diagram_input += "network " + network["meta"]["name"] +" {"
            diagram_input += get_node_NIC_hardware(token, hosts, cluster, network["meta"]["name"])
        except KeyError:
            diagram_input += "network " + network["name"] +" {"
            diagram_input += get_node_NIC_hardware(token, hosts, cluster, network["name"])

        cidr = network["cidr"]
        if str(cidr) != "None":
            diagram_input += 'address = "'+ str(cidr) + '"'
        diagram_input += "}"
    diagram_input += "}"
    return diagram_input

def add_picture_page(filename,page_break=True):
    last = runbook.paragraphs[-1]
    p = last._element
    p.getparent().remove(p)
    p._p = p._element = None

    im = Image.open(filename)
    im.save( filename.replace('.png','.jpg'),'JPEG')

    heading = runbook.add_heading(filename.replace('.png','').replace('screens/',''), level=1)
    heading.alignment = 1
    if "network-layout" in filename:
        runbook.add_picture( filename.replace('.png','.jpg'), width=Inches(4.0))
    else:
        runbook.add_picture( filename.replace('.png','.jpg'), width=Inches(6.5))
    pic = runbook.paragraphs[-1]
    pic.alignment = 1
    if page_break:
        runbook.add_page_break()
    os.remove(filename)

class ForwardServer (SocketServer.ThreadingTCPServer):
    daemon_threads = True
    allow_reuse_address = True

def forward_tunnel(local_port, remote_host, remote_port, transport):
    class SubHander (Handler):
        chain_host = remote_host
        chain_port = remote_port
        ssh_transport = transport
    ForwardServer(('', local_port), SubHander).serve_forever()

class Handler (SocketServer.BaseRequestHandler):
    def handle(self):
        try:
            chan = self.ssh_transport.open_channel('direct-tcpip',(self.chain_host, self.chain_port),self.request.getpeername())
        except Exception as e:
            return
        if chan is None:
            return
        while True:
            r, w, x = select.select([self.request, chan], [], [])
            if self.request in r:
                data = self.request.recv(1024)
                if len(data) == 0:
                    break
                chan.send(data)
            if chan in r:
                data = chan.recv(1024)
                if len(data) == 0:
                    break
                self.request.send(data)

        peername = self.request.getpeername()
        chan.close()
        self.request.close()

def chromedriver():
    # global driver
    driver = webdriver.Chrome(service_args=["--verbose", "--log-path=chromedriver.log"])


def open_tunnels():
    transport = paramiko.Transport((args.host, int(args.ssh_port)))
    transport.connect(username = args.ssh_username,password = args.ssh_password)
    try:
        threads.append(threading.Thread(target=forward_tunnel,args=(8777, "127.0.0.1", 8777, transport),daemon=True).start())
        threads.append(threading.Thread(target=forward_tunnel,args=(8443, "127.0.0.1", 8443, transport),daemon=True).start())
    except:
        sys.exit(0)

# =========================================
# INIT
# =========================================

# Destroy and/or create 'screens/' directory
try:
    shutil.rmtree('screens/')
    os.makedirs('screens/')
except FileNotFoundError:
    os.makedirs('screens/')

# Handle chromedriver dependency
if not cmd_exists('chromedriver'):
    sys.exit('\nYou need chromedriver. Download from here:\nhttps://sites.google.com/a/chromium.org/chromedriver/downloads\n\nAnd install to your path:\nsudo cp chromedriver /usr/local/bin/')

threads = []

open_tunnels()
global driver
# try:
#     threads.append(threading.Thread(target=chromedriver,daemon=True).start())
# except:
#     sys.exit(0)
driver = webdriver.Chrome(service_args=["--verbose", "--log-path=chromedriver.log"])

# Get token & Node informationfrom Fuel API
token = get_token()
# Get Fuel version
version = get_version(token)

cluster = get_cluster(token,args.environment)

nodedata = get_nodes(token)

if args.tests:
    start_ostf(token,args.environment,args.horizon_password,args.horizon_username)

# Init Selenium + chromedriver

driver.set_window_size(1200, 1200)
driver.get('https://localhost:8443')

# Handle Login
try:
    wait_for_page_tag_name('username')
    username = driver.find_element_by_name('username')
    username.send_keys(args.web_username)
    password = driver.find_element_by_name('password')
    password.send_keys(args.web_password)
except NoSuchElementException:
    driver.close()
    sys.exit('Login form not found. Do you have the correct address & port?')
password.send_keys(Keys.RETURN)

# =========================================
# END INIT
# =========================================

# Build cover page
replaces = {
"CUSTOMER": args.customer_name,
"ENV": args.environment_type,
# "RELEASE": entries['COVER']['RELEASE'],
"DATE": time.strftime('%d %B, %Y'),
"AUTHORS": args.deployment_engineer
}

docx_replace('template.docx','cover.docx',replaces)

runbook = Document('cover.docx')

fuel = fuel_info()
gen_access_table()
runbook.add_page_break()

gen_support_table()
runbook.add_page_break()

networkdata = []
gen_nodes_table(args.environment)
runbook.add_page_break()
networkdata = get_network(token,args.environment)
gen_network_layout_table(get_network_layout(),cluster['name'])
runbook.add_page_break()
runbook.add_page_break()

tree = nwdiag.parser.parse_string(create_network_diagram(token, networkdata['networks'], cluster))
diagram = ScreenNodeBuilder.build(tree)
draw = DiagramDraw('PNG', diagram, 'screens/network-layout.png',fontmap=None, antialias=False, nodoctype=True)
draw.draw()
draw.save()
add_picture_page('screens/network-layout.png')

last = runbook.paragraphs[-1]
p = last._element
p.getparent().remove(p)
p._p = p._element = None

if args.tests:
    gen_ostf_table(wait_and_collect_ostf(token,args.environment),cluster['name'])

# Fuel main tabs screenshots
screenshot(None,'Environments','clusters-page')
if '7.0' not in version:
    screenshot('https://localhost:8443' + '/#equipment','Equipment','equipment-page')
screenshot('https://localhost:8443' + '/#releases','Releases','releases-page')
# screenshot('https://localhost:8443' + '/#plugins','Plugins','plugins-page')
#
# # Environments screenshots
#
# # if 'operational' not in e['status']:
# #     break
# c = "https://" + 'localhost' + ':8443' + '/#cluster/'  + str(cluster['id'])
# nodes = []
# for n in nodedata:
#     if n['cluster'] is cluster['id']:
#         nodes.append(n)
#
#
# for i,n in enumerate(nodes):
#     # cluster = str(e['id'])
#     node = str(n['id'])
#     if i == 0:
#         screenshot(c + '/dashboard','Environment: ' + cluster['name'] + ' Dashboard','cluster-page',True)
#         screenshot(c + '/nodes','Environment: ' + cluster['name'] + ' Nodes','nodes-tab',True)
#     screenshot(c + '/nodes/disks/nodes:' +node,'Environment: ' + cluster['name'] + ' Node: ' + n['name'] + ' Disks','edit-node-disks-screen',True )
#     screenshot(c + '/nodes/interfaces/nodes:' +node,'Environment: ' + cluster['name'] + ' Node: ' + n['name'] + ' Interfaces','ifc-container',True )
#
# if '7.0' in version:
#     screenshot(c + '/network','Environment: ' + cluster['name'] + ' Network',None)
# else:
#     driver.get(c + '/network')
#     wait_for_page_tag_class('network-tab')
#     network_elements = driver.find_elements_by_tag_name('a')
#     for i in range(len(network_elements)):
#         element = network_elements[i]
#         if 'subtab-link-' + element.text.lower().replace(' ','_') in element.get_attribute('class'):
#             element.click()
#             driver.execute_script('window.scrollTo(0,0);')
#             time.sleep(.2)
#             screenshot(None,'Environment: ' + cluster['name'] + ' Network ' + element.text,None)
#         if 'Other' in element.text:
#             element.click()
#             time.sleep(.2)
#             screenshot(None,'Environment: ' + cluster['name'] + ' Network Other',None)
#
# driver.get(c + '/settings')
# wait_for_page_tag_class('settings-tab')
# settings_elements = driver.find_elements_by_tag_name('a')
# for i in range(len(settings_elements)-1):
#     element = settings_elements[i]
#     if 'subtab-link-' + element.text.lower().replace(' ','_') in element.get_attribute('class'):
#         element.click()
#         driver.execute_script('window.scrollTo(0,0);')
#         time.sleep(.2)
#         screenshot(None,'Environment: ' + cluster['name'] + ' Settings ' + element.text,None)
#     if 'Kernel parameters' in element.text:
#         element.click()
#         time.sleep(.2)
#         screenshot(None,'Environment: ' + cluster['name'] + ' Settings Kernel Params',None)
#     if 'Repositories' in element.text:
#         element.click()
#         time.sleep(.2)
#         screenshot(None,'Environment: ' + cluster['name'] + ' Settings Repositories',None)
#     if 'Host OS DNS Servers' in element.text:
#         element.click()
#         time.sleep(.2)
#         screenshot(None,'Environment: ' + cluster['name'] + ' Settings DNS',None)
#     if 'Host OS NTP Servers' in element.text:
#         element.click()
#         time.sleep(.2)
#         screenshot(None,'Environment: ' + cluster['name'] + ' Settings NTP',None)
#     if 'Public TLS' in element.text:
#         element.click()
#         time.sleep(.2)
#         screenshot(None,'Environment: ' + cluster['name'] + ' Settings Public TLS',None)
# # driver.get(c + '/healthcheck')
# # wait_for_page_tag_class('healthcheck-controls')
# # screenshot(None,'Environment: ' + cluster['name'] + 'Health Check',None)

# Close browser session
driver.quit()

last = runbook.paragraphs[-1]
p = last._element
p.getparent().remove(p)
p._p = p._element = None

runbook.save(args.filename)

# Cleanup
shutil.rmtree('screens/')
os.remove('cover.docx')
