#!/usr/bin/python3
from docx import Document
from docx.shared import Inches
from docx.enum.style import WD_STYLE
import requests, json

# Open the JSON file
nodeFile = open('nodes.json', 'r')
nodeData = json.loads(nodeFile.read())

# Generate row_count by counting the number of nodes and adding 1 for the header row
row_count = len(nodeData)+1
col_count = 6

# Create (technically open) a new document
doc = Document()

# Add a heading; last digit is heading size
doc.add_heading('Nodes(VMs)',1)

# Create table for the data
table = doc.add_table(row_count, col_count)
table.style = 'TableGrid'

# Generate the header row for the NodesVMS table
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Hostname'
hdr_cells[1].text = 'Role(s)'
hdr_cells[2].text = 'Admin network IP address'
hdr_cells[3].text = 'CPUxCores'
hdr_cells[4].text = 'RAM'
hdr_cells[5].text = 'HDD'

# Modify each additional row using data for the node
for nodeCounter, node in enumerate(nodeData):
	nodeRow = table.rows[nodeCounter+1].cells
	nodeRow[0].text = nodeData[nodeCounter]['hostname']
	nodeRow[1].text = {x+' ' for x in nodeData[nodeCounter]['roles']}
	nodeRow[2].text = nodeData[nodeCounter]['ip']
	nodeRow[3].text = str(nodeData[nodeCounter]['meta']['cpu']['total']) # Total or real?
	nodeRow[4].text = str(int(nodeData[nodeCounter]['meta']['memory']['total']/1048576)) + ' MB' # Total or max cap?
	nodeRow[5].text = {str(x['disk']) + ': ' + str(int(x['size']/1073741824)) + 'GB \n' for x in nodeData[nodeCounter]['meta']['disks']}

doc.save('test.docx')
